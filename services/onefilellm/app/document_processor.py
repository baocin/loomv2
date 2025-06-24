"""Document processor using OneFileLLM."""

import base64
import tempfile
import time
from pathlib import Path

import structlog
from onefilellm import OneFileLLM

from .config import settings
from .models import DocumentTask, ProcessedDocument, ProcessingError

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """Processes documents using OneFileLLM."""

    def __init__(self, kafka_producer, database):
        """Initialize document processor.

        Args:
            kafka_producer: KafkaProducerManager instance
            database: DatabaseManager instance
        """
        self.kafka_producer = kafka_producer
        self.database = database

    async def process_document_task(self, task: DocumentTask) -> None:
        """Process a document task.

        Args:
            task: Document processing task
        """
        start_time = time.time()

        logger.info(
            "Processing document task",
            message_id=task.message_id,
            filename=task.filename,
            document_type=task.document_type,
            file_size=task.file_size,
        )

        try:
            # Decode and save file to temporary location
            with tempfile.TemporaryDirectory() as temp_dir:
                file_path = await self._save_document(task, temp_dir)

                # Process with OneFileLLM
                result = await self._process_with_onefilellm(
                    task, file_path, start_time
                )

                # Send result to Kafka
                await self.kafka_producer.send_processed_document(result)

                # Store in database
                await self.database.store_processed_document(result)

                logger.info(
                    "Document task processed successfully",
                    message_id=task.message_id,
                    processing_time_ms=result.processing_duration_ms,
                    word_count=result.word_count,
                )

        except Exception as e:
            logger.error(
                "Failed to process document task",
                message_id=task.message_id,
                filename=task.filename,
                error=str(e),
            )

            # Send error to Kafka
            error = ProcessingError(
                schema_version=task.schema_version,
                device_id=task.device_id,
                recorded_at=task.recorded_at,
                original_task_id=task.message_id,
                error_type="document_processing_error",
                error_message=str(e),
                error_details={
                    "filename": task.filename,
                    "document_type": task.document_type,
                    "file_size": task.file_size,
                },
                is_retryable=self._is_retryable_error(e),
            )

            await self.kafka_producer.send_processing_error(
                error, settings.topic_document_parsed
            )

    async def _save_document(self, task: DocumentTask, temp_dir: str) -> Path:
        """Save document to temporary file.

        Args:
            task: Document task
            temp_dir: Temporary directory

        Returns:
            Path to saved file
        """
        try:
            # Decode base64 file data
            file_data = base64.b64decode(task.file_data)

            # Create file path with original filename
            file_path = Path(temp_dir) / task.filename

            # Write file data
            with open(file_path, "wb") as f:
                f.write(file_data)

            logger.info(
                "Document saved to temporary file",
                filename=task.filename,
                path=str(file_path),
                size=len(file_data),
            )

            return file_path

        except Exception as e:
            raise Exception(f"Failed to save document: {str(e)}")

    async def _process_with_onefilellm(
        self, task: DocumentTask, file_path: Path, start_time: float
    ) -> ProcessedDocument:
        """Process document with OneFileLLM.

        Args:
            task: Original document task
            file_path: Path to saved document
            start_time: Processing start time

        Returns:
            Processed document result
        """
        logger.info("Processing document with OneFileLLM", file_path=str(file_path))

        try:
            # Determine extraction method based on file type
            extraction_method = self._get_extraction_method(
                task.filename, task.content_type
            )

            # Extract text content
            extracted_text = await self._extract_text_content(
                file_path, extraction_method
            )

            # Calculate text statistics
            word_count = len(extracted_text.split())
            character_count = len(extracted_text)

            # Detect language (simple heuristic)
            language_detected = self._detect_language(extracted_text)

            # Generate summary
            content_summary = self._generate_summary(extracted_text, task.filename)

            # Get page count for PDFs
            page_count = self._get_page_count(file_path, task.content_type)

            # Calculate processing duration
            processing_duration_ms = int((time.time() - start_time) * 1000)

            # Create result
            result = ProcessedDocument(
                schema_version=task.schema_version,
                device_id=task.device_id,
                recorded_at=task.recorded_at,
                original_filename=task.filename,
                document_type=task.document_type,
                content_type=task.content_type,
                extracted_text=extracted_text,
                content_summary=content_summary,
                original_size_bytes=task.file_size,
                processing_duration_ms=processing_duration_ms,
                extraction_method=extraction_method,
                language_detected=language_detected,
                page_count=page_count,
                word_count=word_count,
                character_count=character_count,
                extraction_metadata={
                    "original_metadata": task.metadata,
                    "extract_options": task.extract_options,
                },
                onefilellm_version=self._get_onefilellm_version(),
            )

            return result

        except Exception as e:
            raise Exception(f"Document processing failed: {str(e)}")

    def _get_extraction_method(self, filename: str, content_type: str) -> str:
        """Determine extraction method based on file type.

        Args:
            filename: Original filename
            content_type: MIME type

        Returns:
            Extraction method name
        """
        filename_lower = filename.lower()

        if filename_lower.endswith(".pdf") or "pdf" in content_type:
            return "pdf_extraction"
        elif filename_lower.endswith((".docx", ".doc")) or "word" in content_type:
            return "docx_extraction"
        elif filename_lower.endswith((".txt", ".md", ".rst")):
            return "text_extraction"
        elif filename_lower.endswith((".html", ".htm")) or "html" in content_type:
            return "html_extraction"
        else:
            return "onefilellm_extraction"

    async def _extract_text_content(
        self, file_path: Path, extraction_method: str
    ) -> str:
        """Extract text content from document.

        Args:
            file_path: Path to document
            extraction_method: Method to use for extraction

        Returns:
            Extracted text content
        """
        try:
            if extraction_method == "pdf_extraction":
                return await self._extract_pdf_text(file_path)
            elif extraction_method == "docx_extraction":
                return await self._extract_docx_text(file_path)
            elif extraction_method == "html_extraction":
                return await self._extract_html_text(file_path)
            elif extraction_method == "text_extraction":
                return await self._extract_plain_text(file_path)
            else:
                # Use OneFileLLM for unknown file types
                return await self._extract_with_onefilellm(file_path)

        except Exception as e:
            logger.warning(
                "Primary extraction failed, falling back to OneFileLLM",
                method=extraction_method,
                error=str(e),
            )
            # Fallback to OneFileLLM
            return await self._extract_with_onefilellm(file_path)

    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using pypdf.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        import pypdf

        text_parts = []

        with open(file_path, "rb") as file:
            reader = pypdf.PdfReader(file)

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)

        return "\n\n".join(text_parts)

    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        import docx

        doc = docx.Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        return "\n\n".join(text_parts)

    async def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML using beautifulsoup4.

        Args:
            file_path: Path to HTML file

        Returns:
            Extracted text
        """
        from bs4 import BeautifulSoup

        with open(file_path, encoding="utf-8") as file:
            soup = BeautifulSoup(file.read(), "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            return text

    async def _extract_plain_text(self, file_path: Path) -> str:
        """Extract plain text from text file.

        Args:
            file_path: Path to text file

        Returns:
            File content
        """
        import chardet

        # Detect encoding
        with open(file_path, "rb") as file:
            raw_data = file.read()
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result.get("encoding", "utf-8")

        # Read with detected encoding
        with open(file_path, encoding=encoding) as file:
            return file.read()

    async def _extract_with_onefilellm(self, file_path: Path) -> str:
        """Extract content using OneFileLLM.

        Args:
            file_path: Path to file

        Returns:
            Extracted content
        """
        import asyncio

        # Create temporary directory containing just this file
        parent_dir = file_path.parent

        # Initialize OneFileLLM
        one_file_llm = OneFileLLM()

        # Run OneFileLLM processing in thread pool
        def _process():
            return one_file_llm.fit(
                input_directory=str(parent_dir),
                include_patterns=[file_path.name],  # Include only this file
                exclude_patterns=[],
            )

        loop = asyncio.get_event_loop()
        result = await asyncio.wait_for(
            loop.run_in_executor(None, _process),
            timeout=settings.processing_timeout,
        )

        return result.content

    def _detect_language(self, text: str) -> str | None:
        """Simple language detection.

        Args:
            text: Text content

        Returns:
            Detected language code or None
        """
        # Simple heuristic - just return English for now
        # In a real implementation, you might use langdetect or similar
        if text.strip():
            return "en"
        return None

    def _generate_summary(self, text: str, filename: str) -> str:
        """Generate a brief summary of the document.

        Args:
            text: Document text
            filename: Original filename

        Returns:
            Brief summary
        """
        word_count = len(text.split())
        char_count = len(text)

        # Get first few sentences as a preview
        sentences = text.split(". ")[:3]
        preview = ". ".join(sentences)
        if len(preview) > 200:
            preview = preview[:200] + "..."

        summary_parts = [
            f"Document '{filename}' contains {word_count:,} words ({char_count:,} characters)"
        ]

        if preview.strip():
            summary_parts.append(f"Content preview: {preview}")

        return ". ".join(summary_parts) + "."

    def _get_page_count(self, file_path: Path, content_type: str) -> int | None:
        """Get page count for documents that support it.

        Args:
            file_path: Path to document
            content_type: MIME type

        Returns:
            Page count or None
        """
        try:
            if "pdf" in content_type or file_path.suffix.lower() == ".pdf":
                import pypdf

                with open(file_path, "rb") as file:
                    reader = pypdf.PdfReader(file)
                    return len(reader.pages)
        except Exception:
            pass

        return None

    def _get_onefilellm_version(self) -> str:
        """Get OneFileLLM version.

        Returns:
            OneFileLLM version string
        """
        try:
            import onefilellm

            return getattr(onefilellm, "__version__", "unknown")
        except Exception:
            return "unknown"

    def _is_retryable_error(self, error: Exception) -> bool:
        """Determine if an error is retryable.

        Args:
            error: Exception that occurred

        Returns:
            True if error is retryable
        """
        error_str = str(error).lower()

        # Non-retryable errors
        non_retryable_indicators = [
            "invalid",
            "corrupt",
            "unsupported",
            "malformed",
            "encoding",
        ]

        # Retryable errors
        retryable_indicators = [
            "timeout",
            "temporary",
            "memory",
        ]

        for indicator in non_retryable_indicators:
            if indicator in error_str:
                return False

        for indicator in retryable_indicators:
            if indicator in error_str:
                return True

        # Default to retryable for unknown errors
        return True
